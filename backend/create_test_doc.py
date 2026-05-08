import docx
import os

def create_test_docx():
    doc = docx.Document()
    doc.add_heading('AstroDynamics Corp - Internal Operations Manual', 0)
    
    doc.add_heading('1. Company Overview', level=1)
    doc.add_paragraph('AstroDynamics Corp was founded by Dr. Elena Quasar in 2021. The company is headquartered in Neo-Tokyo and specializes in deep-space propulsion systems.')
    
    doc.add_heading('2. The Starlight Engine', level=2)
    doc.add_paragraph('The Starlight Engine is our flagship product. It uses a proprietary Ion-Fusion hybrid technology that allows for continuous acceleration up to 0.1c (10% the speed of light).')
    doc.add_paragraph('Official Launch Date: October 14, 2025.')
    
    doc.add_heading('3. Known Issues and Maintenance', level=1)
    doc.add_paragraph('Engineers should be aware of the following documented issues:')
    doc.add_paragraph('• Fuel Leakage: A minor fuel leak may occur in zero-G environments when the primary valve is at 45% capacity.')
    doc.add_paragraph('• Overheating: Some units have reported overheating during atmospheric exit if the cooling cycle is not engaged 5 minutes prior to takeoff.')
    
    doc.add_heading('4. Contact Information', level=1)
    doc.add_paragraph('For emergency support, contact Chief Engineer Marcus Sol at m.sol@astrodynamics.internal or call the 24/7 hotline at 555-STAR-99.')

def create_product_roadmap():
    doc = docx.Document()
    doc.add_heading('EcoTrack AI - Project Roadmap 2025-2026', 0)
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph('EcoTrack AI is a revolutionary platform designed to track carbon emissions in real-time using satellite imagery and IoT sensors.')
    
    doc.add_heading('2. Project Phases', level=1)
    doc.add_paragraph('• Phase 1: Data Collection & Sensor Integration. Timeline: Q3 2025.')
    doc.add_paragraph('• Phase 2: AI Model Training & Validation. Timeline: Q4 2025.')
    doc.add_paragraph('• Phase 3: Global Public Deployment. Timeline: Q1 2026.')
    
    doc.add_heading('3. Financials', level=1)
    doc.add_paragraph('Total Allocated Budget: $2.5 Million USD.')
    doc.add_paragraph('Current Spending: $450,000 USD.')
    
    doc.add_heading('4. Leadership Team', level=1)
    doc.add_paragraph('Project Lead: Sarah Green.')
    doc.add_paragraph('Technical Director: Dr. Alan Turing (AI Division).')

    file_path = "product_roadmap.docx"
    doc.save(file_path)
    print(f"File created successfully at: {os.path.abspath(file_path)}")

def create_resume():
    doc = docx.Document()
    doc.add_heading('John Doe - Senior Software Engineer', 0)
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Results-oriented Software Engineer with 8 years of experience in Python, React, and Cloud Architecture.')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('TechFlow Solutions | Senior Developer | 2020 - Present')
    doc.add_paragraph('• Led a team of 5 to build a real-time analytics dashboard.')
    doc.add_paragraph('• Reduced server costs by 30% using AWS Lambda.')
    
    doc.add_paragraph('DataStream Inc | Junior Developer | 2017 - 2020')
    doc.add_paragraph('• Maintained legacy Java applications and migrated them to Node.js.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science | Stanford University | 2017')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Languages: Python, JavaScript, SQL, C++.')
    doc.add_paragraph('Frameworks: React, FastAPI, Next.js, LangChain.')

    file_path = "john_doe_resume.docx"
    doc.save(file_path)
    print(f"File created successfully at: {os.path.abspath(file_path)}")

if __name__ == "__main__":
    create_test_docx()
    create_product_roadmap()
    create_resume()
