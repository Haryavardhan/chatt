import os
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import docx
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="RAG Debugging Tool")

# CORS setup
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Constants
DATA_DIR = "data"
VECTOR_DB_DIR = "vector_db"
DOC_PATH = os.path.join(DATA_DIR, "debugging.docx")
MODEL_NAME = "mixtral-8x7b-32768"

# Ensure directories exist
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(VECTOR_DB_DIR, exist_ok=True)

# Global variables
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
vector_store = None
llm = None

def create_sample_docx():
    if not os.path.exists(DOC_PATH):
        doc = docx.Document()
        doc.add_heading('Comprehensive Project Debugging & Architecture Guide', 0)
        
        doc.add_heading('1. Backend: FastAPI & Python Errors', level=1)
        doc.add_paragraph('ModuleNotFoundError: This is the most common error. It occurs when a dependency is missing. Fix: Run `pip install -r requirements.txt`.')
        doc.add_paragraph('OperationalError: Usually related to database connections. Check if your DB server is up and the credentials in .env are correct.')
        doc.add_paragraph('ValidationError: Pydantic raises this when request data does not match the defined model schema.')
        
        doc.add_heading('2. Frontend: Next.js & React Errors', level=1)
        doc.add_paragraph('Hydration Failed: This happens when the server-rendered HTML differs from the client-rendered HTML.')
        doc.add_paragraph('AxiosError: Network errors during API calls.')

        doc.add_heading('3. Project Architecture Overview', level=1)
        doc.add_paragraph('Frontend: Next.js 14+, Tailwind CSS, Framer Motion.')
        doc.add_paragraph('Backend: FastAPI, LangChain, FAISS.')
        
        doc.save(DOC_PATH)
        return True
    return False

def init_rag():
    global vector_store, llm
    
    if not os.path.exists(DOC_PATH):
        create_sample_docx()
    
    # Load and process docx
    doc = docx.Document(DOC_PATH)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_text(full_text)
    docs = [Document(page_content=chunk) for chunk in chunks]
    
    # Create vector store
    vector_store = FAISS.from_documents(docs, embeddings)
    vector_store.save_local(VECTOR_DB_DIR)
    
    # Init LLM
    llm = ChatGroq(
        temperature=0,
        model_name=MODEL_NAME,
        groq_api_key=os.getenv("GROQ_API_KEY")
    )

@app.on_event("startup")
async def startup_event():
    try:
        init_rag()
    except Exception as e:
        print(f"Error initializing RAG: {e}")

class QueryRequest(BaseModel):
    question: str

@app.post("/ask")
async def ask_question(request: QueryRequest):
    if not vector_store or not llm:
        raise HTTPException(status_code=500, detail="System not initialized")
    
    try:
        # 1. Retrieve context
        context_docs = vector_store.similarity_search(request.question, k=3)
        context_text = "\n\n".join([doc.page_content for doc in context_docs])
        
        # 2. Build prompt
        prompt = f"""You are a debugging assistant. Use the following context to answer the question.
If the answer isn't in the context, use your general knowledge but mention it's not in the doc.

Context:
{context_text}

Question: {request.question}
Answer:"""

        # 3. Call LLM
        response = llm.invoke(prompt)
        return {"answer": response.content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are allowed")
    
    with open(DOC_PATH, "wb") as f:
        f.write(await file.read())
    
    init_rag()
    return {"message": "Document uploaded and indexed successfully"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
